# clavis_event_inventory/inventory/utils.py

import io
from datetime import date
from django.http import HttpResponse
from django.conf import settings # Ensure settings is imported
import os # Ensure os is imported

# Excel Export
import openpyxl
from openpyxl.styles import Font
# from openpyxl.drawing.image import Image as OpenpyxlImage # For Excel images, if added later

# PDF Export
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image # Ensure Image is imported
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from PIL import Image as PILImage, ImageOps

# Word Export
from docx import Document
from docx.shared import Inches, Pt
# from docx.shared import Cm # For Word images, if added later

# generate barcode start
import barcode
from barcode.writer import ImageWriter
from io import BytesIO
import base64

def generate_barcode_base64(sku):
    if not sku:
        return None
    CODE128 = barcode.get_barcode_class('code128')
    buffer = BytesIO()
    CODE128(sku, writer=ImageWriter()).write(buffer)
    return base64.b64encode(buffer.getvalue()).decode('utf-8')

# generate barcode end

def generate_master_inventory_excel(items):
    """Generates an Excel HttpResponse for the master inventory list."""
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="clavis_master_inventory_{date.today()}.xlsx"'

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = 'Master Inventory'

    # Headers (Depth instead of Length)
    headers = [
        "SKU", "Item Name", "Category", "Description", "Location",
        "Initial Qty", "Available Qty", "Purchase Price (BHD)", "Rent Price/Day (BHD)",
        "Supplier", "Depth", "Width", "Height", "Unit", "Created", "Updated"
    ]
    sheet.append(headers)

    for cell in sheet[1]:
        cell.font = Font(bold=True)

    # Data Rows (using item.depth)
    for item in items:
        sheet.append([
            item.sku,
            item.name,
            item.category.name if item.category else '-',
            item.description,
            item.storage_location,
            item.initial_quantity,
            item.available_quantity,
            item.purchase_price,
            item.rent_price_per_day,
            item.supplier.name if item.supplier else '-',
            item.depth, # Using depth
            item.width,
            item.height,
            item.get_dimension_unit_display(),
            item.created_at.strftime('%Y-%m-%d') if item.created_at else '-',
            item.updated_at.strftime('%Y-%m-%d') if item.updated_at else '-',
        ])

    # Adjust column widths
    for col_idx, column_cells in enumerate(sheet.columns):
        # Skip image column if we add it later to Excel, as width is tricky
        # if headers[col_idx] == "Image":
        #     sheet.column_dimensions[column_cells[0].column_letter].width = 15 # Example fixed width for image
        #     continue
        
        max_length = 0
        column_letter = column_cells[0].column_letter
        for cell in column_cells:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except: pass
        adjusted_width = (max_length + 2)
        if column_letter in ['H', 'I'] and adjusted_width < 15: adjusted_width = 15
        if column_letter in ['K','L','M'] and adjusted_width < 8: adjusted_width = 8
        sheet.column_dimensions[column_letter].width = adjusted_width

    workbook.save(response)
    return response


def generate_master_inventory_pdf(items):
    """Generates a PDF HttpResponse for the master inventory list with images."""
    buffer = io.BytesIO()
    # Consider portrait if too many columns make landscape too cramped with images
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), rightMargin=0.5*inch, leftMargin=0.5*inch, topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    story = []

    # Define a paragraph style for table cells for better text wrapping
    table_cell_style = ParagraphStyle('TableCell', parent=styles['Normal'], fontSize=7, leading=8)


    title = f"Clavis Master Inventory Report - {date.today()}"
    story.append(Paragraph(title, styles['h1']))
    story.append(Spacer(1, 0.2*inch))

    # ADDED "Image" to headers
    headers = ["Image", "SKU", "Item Name", "Category", "Loc", "Init", "Avail", "Purch", "Rent", "D", "W", "H", "Unit", "Supplier"]
    data = [headers]

    for item in items:
        purch_price = f"{item.purchase_price} BHD" if item.purchase_price is not None else '-'
        rent_price = f"{item.rent_price_per_day} BHD" if item.rent_price_per_day is not None else '-'
        
        # --- Image Handling ---
        item_image_el = Paragraph("(No Image)", table_cell_style) # Placeholder
        if item.image1 and hasattr(item.image1, 'path'):
            try:
                if os.path.exists(item.image1.path):
                    with PILImage.open(item.image1.path) as pil_img:
                        pil_img = ImageOps.exif_transpose(pil_img)  # Auto-rotate if needed
                        pil_img = pil_img.convert("RGB")  # Ensure compatible mode for PDF

                        img_buffer = io.BytesIO()
                        pil_img.save(img_buffer, format='JPEG')
                        img_buffer.seek(0)

                        img = Image(img_buffer, width=0.5*inch, height=0.5*inch)
                        img.hAlign = 'CENTER'
                        item_image_el = img
            except Exception as e:
                print(f"Error loading image {item.image1.path} for PDF: {e}")
        # --- End Image Handling ---

        data.append([
            item_image_el, # ADDED Image element or placeholder
            Paragraph(item.sku or '-', table_cell_style), 
            Paragraph(item.name or '-', table_cell_style), 
            Paragraph(item.category.name if item.category else '-', table_cell_style),
            Paragraph(item.storage_location or '-', table_cell_style), 
            item.initial_quantity, 
            item.available_quantity,
            Paragraph(purch_price, table_cell_style), 
            Paragraph(rent_price, table_cell_style),
            item.depth or '-', 
            item.width or '-', 
            item.height or '-',
            Paragraph(item.get_dimension_unit_display(), table_cell_style), 
            Paragraph(item.supplier.name if item.supplier else '-', table_cell_style),
        ])

    # ADDED width for Image column, adjusted others
    col_widths = [0.7*inch, 0.7*inch, 1.2*inch, 0.7*inch, 0.7*inch, 0.4*inch, 0.4*inch, 0.7*inch, 0.7*inch, 0.4*inch, 0.4*inch, 0.4*inch, 0.4*inch, 0.8*inch ]

    if len(data) > 1 :
        table = Table(data, colWidths=col_widths, rowHeights=0.6*inch) # Set a default row height for images
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey), 
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'), 
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'), # Vertically align content in cells
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8), 
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black), 
            ('FONTSIZE', (0, 0), (-1, -1), 7), # Keep font small for data
            ('ALIGN', (1, 1), (2, -1), 'LEFT'), # Align SKU, Item Name left
            ('ALIGN', (3, 1), (4, -1), 'LEFT'), # Align Category, Loc left
            ('ALIGN', (7, 1), (8, -1), 'RIGHT'), # Align prices right
            ('ALIGN', (13, 1), (13, -1), 'LEFT'), # Align Supplier left
        ]))
        story.append(table)
    else:
        story.append(Paragraph("No inventory items found.", styles['Normal']))

    doc.build(story)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="clavis_master_inventory_{date.today()}.pdf"'
    return response


def generate_master_inventory_docx(items):
    """Generates a Word (.docx) HttpResponse for the master inventory list."""
    # Note: Adding images to DOCX tables cell by cell can be complex and might require
    # more advanced manipulation or a different approach for optimal layout.
    # This version will NOT include images for simplicity, but can be extended.
    document = Document()
    document.add_heading(f"Clavis Master Inventory Report - {date.today()}", level=1)

    headers = [ "SKU", "Item Name", "Category", "Location", "Initial", "Avail", "Purch Price", "Rent/Day", "D", "W", "H", "Unit", "Supplier" ]

    if items:
        table = document.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers): 
            run = hdr_cells[i].paragraphs[0].add_run(header)
            run.font.bold = True; run.font.size = Pt(8)

        for item in items:
            row_cells = table.add_row().cells
            purch_price = f"{item.purchase_price} BHD" if item.purchase_price is not None else '-'
            rent_price = f"{item.rent_price_per_day} BHD" if item.rent_price_per_day is not None else '-'
            
            row_cells[0].text = item.sku or '-'; 
            row_cells[1].text = item.name or '-';
            row_cells[2].text = item.category.name if item.category else '-'; 
            row_cells[3].text = item.storage_location or '-';
            row_cells[4].text = str(item.initial_quantity); 
            row_cells[5].text = str(item.available_quantity);
            row_cells[6].text = purch_price; 
            row_cells[7].text = rent_price;
            row_cells[8].text = str(item.depth or '-'); 
            row_cells[9].text = str(item.width or '-');
            row_cells[10].text = str(item.height or '-'); 
            row_cells[11].text = item.get_dimension_unit_display();
            row_cells[12].text = item.supplier.name if item.supplier else '-'
            for cell in row_cells: 
                if cell.paragraphs: # Ensure paragraph exists
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(8)
                else: # Add a paragraph if none exists
                    p = cell.add_paragraph()
                    p.add_run().font.size = Pt(8)


    else:
         document.add_paragraph("No inventory items found.")

    buffer = io.BytesIO(); document.save(buffer); buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="clavis_master_inventory_{date.today()}.docx"'
    return response
