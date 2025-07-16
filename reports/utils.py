# clavis_event_inventory/reports/utils.py

import io
from django.http import HttpResponse
from django.utils import timezone
import calendar
from django.conf import settings
import os
from datetime import date

# Excel Export
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

# PDF Export
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

# Word Export
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement

# Import models
from clients.models import Client
from bookings.models import Event, Rental # For type checking and status choices

# --- Company Details (Consider moving to a central config or settings if used elsewhere) ---
COMPANY_NAME_FOR_PDF_FALLBACK = "Clavis Events & Promotions W.L.L"
COMPANY_ADDRESS_PDF = "Office 123, Building 456, Road 789<br/>Manama, Kingdom of Bahrain" 
COMPANY_CONTACT_PDF = "Tel: +973 1700 0000 | Email: info@clavisevents.com"
COMPANY_CR_VAT_PDF = "CR: XXXXXX-X | VAT: XXXXXXXXXXXXXXX"

def get_month_year_str(year, month):
    if str(year) == "all" and str(month) == "all":
        return "All Years, All Months"
    elif str(year) == "all":
        try:
            return f"All Years - {calendar.month_name[int(month)]}"
        except:
            return "All Years - Invalid Month"
    elif str(month) == "all":
        return f"{year} - All Months"
    else:
        try:
            return f"{calendar.month_name[int(month)]} {year}"
        except:
            return f"{month} {year}"

# --- PDF STYLES ---
def get_report_pdf_styles():
    styles = getSampleStyleSheet()
    base_font = 'Helvetica'
    base_font_bold = 'Helvetica-Bold'

    styles.add(ParagraphStyle(name='ReportMainTitle', parent=styles['h1'], fontName=base_font_bold, fontSize=18, leading=22, alignment=TA_CENTER, spaceAfter=6, textColor=colors.HexColor("#1A237E")))
    styles.add(ParagraphStyle(name='ReportPeriodSubtitle', parent=styles['Normal'], alignment=TA_CENTER, fontSize=11, textColor=colors.dimgrey, spaceAfter=18, fontName=base_font))
    styles.add(ParagraphStyle(name='ReportSectionTitle', parent=styles['h2'], fontName=base_font_bold, fontSize=12, leading=14, textColor=colors.HexColor("#2C3E50"), spaceBefore=10, spaceAfter=4, borderPadding=2))
    styles.add(ParagraphStyle(name='ReportTableHeader', parent=styles['Normal'], fontName=base_font_bold, fontSize=8, alignment=TA_CENTER, textColor=colors.whitesmoke, leading=10))
    styles.add(ParagraphStyle(name='ReportTableCell', parent=styles['Normal'], fontName=base_font, fontSize=7.5, leading=9, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name='ReportTableCellRight', parent=styles['ReportTableCell'], alignment=TA_RIGHT))
    styles.add(ParagraphStyle(name='ReportTableCellCenter', parent=styles['ReportTableCell'], alignment=TA_CENTER))
    styles.add(ParagraphStyle(name='NoDataMessage', parent=styles['Normal'], fontSize=9, textColor=colors.grey, alignment=TA_CENTER, spaceBefore=6, spaceAfter=6))
    return styles

def get_base_table_style(header_bg_color=colors.HexColor("#4A5568"), grid_color=colors.darkgrey, even_row_bg_color=colors.HexColor("#F0F8FF"), odd_row_bg_color=colors.white):
    style_cmds = [
        ('BACKGROUND', (0,0), (-1,0), header_bg_color),
        ('TEXTCOLOR',(0,0),(-1,0), colors.whitesmoke),
        ('ALIGN', (0,0), (-1,0), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,0), 7.5),
        ('FONTSIZE', (0,1), (-1,-1), 7),
        ('GRID', (0,0), (-1,-1), 0.5, grid_color),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
        ('LEFTPADDING', (0,0), (-1,-1), 4),
        ('RIGHTPADDING', (0,0), (-1,-1), 4),
    ]
    return style_cmds

def apply_zebra_striping(table_style_list, data_rows_count, even_row_bg_color=colors.HexColor("#F0F8FF"), odd_row_bg_color=colors.white):
    for i in range(1, data_rows_count + 1):
        bg_color = even_row_bg_color if i % 2 == 0 else odd_row_bg_color
        table_style_list.append(('BACKGROUND', (0, i), (-1, i), bg_color))
    return table_style_list


# --- Excel Generation ---
# MODIFIED: Added logistics_services parameter
def generate_monthly_summary_excel(year, month, regular_events, logistics_services, rentals, clients_summary, item_usage_summary):
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    if str(month).isdigit():
        month_int = int(month)
        month_name_short = calendar.month_abbr[month_int]
    else:
        month_int = None
        month_name_short = "All"

    month_year_str = get_month_year_str(year, month)
    response['Content-Disposition'] = f'attachment; filename="Clavis_Monthly_Summary_{month_name_short}_{year}.xlsx"'

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = f"Summary {month_name_short} {year}"

    # Styles
    title_font = Font(name='Calibri', size=16, bold=True, color="FF000080")
    subtitle_font = Font(name='Calibri', size=12, italic=True, color="FF4F4F4F")
    section_header_font = Font(name='Calibri', size=14, bold=True, color="FF1F4E78")
    table_header_font = Font(name='Calibri', size=10, bold=True, color="FFFFFFFF")
    table_header_fill = PatternFill(start_color="FF4A90E2", end_color="FF4A90E2", fill_type="solid")
    center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
    right_alignment = Alignment(horizontal='right', vertical='center', wrap_text=True)
    thin_border_side = Side(style='thin')
    thin_border = Border(left=thin_border_side, right=thin_border_side, top=thin_border_side, bottom=thin_border_side)

    current_row = 1
    sheet.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=7)
    main_title_cell = sheet.cell(row=current_row, column=1, value="Clavis Monthly Summary Report")
    main_title_cell.font = title_font
    main_title_cell.alignment = center_alignment
    current_row += 1
    
    sheet.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=7)
    period_cell = sheet.cell(row=current_row, column=1, value=f"Period: {month_year_str}")
    period_cell.font = subtitle_font
    period_cell.alignment = center_alignment
    current_row += 2

    def add_table_section_excel(title, headers, data_queryset, data_mapper_func, current_r):
        if headers:
            sheet.merge_cells(start_row=current_r, start_column=1, end_row=current_r, end_column=len(headers))
        else:
            sheet.merge_cells(start_row=current_r, start_column=1, end_row=current_r, end_column=1)

        title_cell = sheet.cell(row=current_r, column=1, value=title)
        title_cell.font = section_header_font
        title_cell.alignment = left_alignment
        current_r += 1
        
        has_data = False
        if hasattr(data_queryset, 'exists'): # For QuerySets
            has_data = data_queryset.exists()
        elif isinstance(data_queryset, list): # For lists
            has_data = bool(data_queryset)

        if has_data:
            for col_idx, header_title in enumerate(headers, 1):
                cell = sheet.cell(row=current_r, column=col_idx, value=header_title)
                cell.font = table_header_font
                cell.fill = table_header_fill
                cell.alignment = center_alignment
                cell.border = thin_border
            current_r +=1
            for item_obj in data_queryset:
                row_values = data_mapper_func(item_obj)
                sheet.append(row_values) 
                data_row_index = sheet.max_row 
                for col_idx_val, val in enumerate(row_values, 1):
                    data_cell = sheet.cell(row=data_row_index, column=col_idx_val)
                    data_cell.border = thin_border
                    is_numeric_header = headers[col_idx_val-1] not in ["Ref #", "Client", "Event Name", "Location", "Service Title", "Primary Site", "Status", "Company", "Email", "Registered On", "Item Name", "SKU"]
                    if isinstance(val, (int, float)) and is_numeric_header:
                        data_cell.alignment = right_alignment
                    else:
                        data_cell.alignment = left_alignment
            current_r = sheet.max_row + 1 
        else:
            sheet.cell(row=current_r, column=1, value=f"No {title.split(' (')[0].lower().replace('handled','').replace('activity','').replace('registered','').strip()} for this period.")
            current_r += 1
        return current_r + 1

    # Regular Events Section
    current_row = add_table_section_excel(
        f"Events Handled ({regular_events.count()})",
        ["Ref #", "Client", "Event Name", "Location", "Start Date", "End Date", "Status"],
        regular_events,
        lambda e: [
            e.reference_number or '-', 
            f"{e.client.name} ({e.client.company_name})" if e.client and e.client.company_name else (str(e.client) if e.client else '-'),
            e.event_name or '-',
            e.event_location or '-', 
            timezone.localtime(e.start_date).strftime('%Y-%m-%d %H:%M') if e.start_date else '-',
            timezone.localtime(e.end_date).strftime('%Y-%m-%d %H:%M') if e.end_date else '-',
            e.get_status_display() or '-'
        ],
        current_row
    )

    # Logistics Services Section
    current_row = add_table_section_excel(
        f"Logistics Services Handled ({logistics_services.count()})",
        ["Ref #", "Client", "Service Title", "Primary Site", "Start Date", "End Date", "Status"],
        logistics_services,
        lambda s: [
            s.reference_number or '-', 
            f"{s.client.name} ({s.client.company_name})" if s.client and s.client.company_name else (str(s.client) if s.client else '-'),
            s.event_name or '-', 
            s.event_location or '-', 
            timezone.localtime(s.start_date).strftime('%Y-%m-%d %H:%M') if s.start_date else '-',
            timezone.localtime(s.end_date).strftime('%Y-%m-%d %H:%M') if s.end_date else '-',
            s.get_status_display() or '-'
        ],
        current_row
    )

    # Rentals Section
    current_row = add_table_section_excel(
        f"Rentals Handled ({rentals.count()})",
        ["Ref #", "Client", "Pickup Date", "Return Date", "Status"],
        rentals,
        lambda r: [
            r.reference_number or '-', 
            f"{r.client.name} ({r.client.company_name})" if r.client and r.client.company_name else (str(r.client) if r.client else '-'),
            timezone.localtime(r.start_date).strftime('%Y-%m-%d %H:%M') if r.start_date else '-',
            timezone.localtime(r.end_date).strftime('%Y-%m-%d %H:%M') if r.end_date else '-',
            r.get_status_display() or '-'
        ],
        current_row
    )

    # Clients Summary Section
    new_clients = clients_summary.get('new_clients')
    current_row = add_table_section_excel(
        f"New Clients Registered ({clients_summary.get('new_clients_count', 0)})",
        ["Client Name", "Company", "Email", "Registered On"],
        new_clients,
        lambda c: [
            c.name or '-', c.company_name or '-', 
            c.email or '-', 
            timezone.localtime(c.created_at).strftime('%Y-%m-%d') if c.created_at else '-'
        ],
        current_row
    )

    # Item Usage Summary
    current_row = add_table_section_excel(
        f"Item Usage Summary ({len(item_usage_summary)} item types used)",
        ["Item Name", "SKU", "Times Used", "Total Qty Used"],
        item_usage_summary, 
        lambda i: [
            i.get('item__name', '-'), i.get('item__sku', '-'),
            i.get('times_used', 0), i.get('total_quantity_used', 0)
        ],
        current_row
    )
    
    # Auto-size columns for all data
    for column_cells_group in sheet.columns:
        max_len = 0
        column_letter_excel = openpyxl.utils.get_column_letter(column_cells_group[0].column)
        for cell_item in column_cells_group:
            try:
                if cell_item.value:
                    cell_value_str = str(cell_item.value)
                    # Handle potential multiline content by splitting and taking max line length
                    current_cell_max_len = max(len(line) for line in cell_value_str.split('\n'))
                    if current_cell_max_len > max_len:
                        max_len = current_cell_max_len
            except:
                pass
        adjusted_width = (max_len + 2) if max_len > 0 else 12 
        sheet.column_dimensions[column_letter_excel].width = adjusted_width

    workbook.save(response)
    return response

# --- PDF Generation ---
# MODIFIED: Added logistics_services parameter
def generate_monthly_summary_pdf(year, month, regular_events, logistics_services, rentals, clients_summary, item_usage_summary):
    buffer = io.BytesIO()
    # month_int = int(month)
    # month_year_str = get_month_year_str(year, month_int) # Corrected variable name
    month_year_str = get_month_year_str(year, month)
    
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                            rightMargin=0.5*inch, leftMargin=0.5*inch, 
                            topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = get_report_pdf_styles() # Use report-specific styles
    story = []

    # --- Logo ---
    logo_path = os.path.join(settings.STATICFILES_DIRS[0] if settings.STATICFILES_DIRS else settings.STATIC_ROOT or '', 'images', 'clavis_logo.png')
    logo_el = None
    if os.path.exists(logo_path):
        try:
            logo = Image(logo_path, width=1.5*inch, height=0.6*inch, kind='proportional') 
            logo_el = logo
        except Exception as e:
            print(f"--- ERROR: Error loading logo image for PDF: {e}")
    
    report_title_style = styles['ReportMainTitle']
    report_period_style = styles['ReportPeriodSubtitle']
    h2_style_pdf = styles['ReportSectionTitle']
    table_cell_style_pdf = styles['ReportTableCell']
    table_header_cell_style_pdf = styles['ReportTableHeader']
    no_data_style = styles['NoDataMessage']

    # --- Header Table ---
    header_left_elements = [
        Paragraph("Clavis Monthly Summary Report", report_title_style),
        Paragraph(f"Period: {month_year_str}", report_period_style)
    ]
    if logo_el:
        header_table_data = [[header_left_elements, logo_el]]
        header_table = Table(header_table_data, colWidths=[5.5*inch, 1.5*inch]) 
        header_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'), ('ALIGN', (0,0), (0,0), 'LEFT'), 
            ('ALIGN', (1,0), (1,0), 'RIGHT'), ('LEFTPADDING', (0,0), (0,0), 0), 
            ('RIGHTPADDING', (1,0), (1,0), 0),
        ]))
        story.append(header_table)
    else:
        story.extend(header_left_elements)
    story.append(Spacer(1, 0.1*inch))
    
    current_page_elements = [] # To track elements for potential page break

    def add_pdf_table_section(title, headers, queryset, mapper_func, col_widths, story_list, styles_dict):
        story_list.append(Paragraph(title, styles_dict['ReportSectionTitle']))
        if queryset.exists() if hasattr(queryset, 'exists') else queryset:
            table_headers = [Paragraph(h, styles_dict['ReportTableHeader']) for h in headers]
            table_data = [table_headers]
            for item_obj in queryset:
                table_data.append([Paragraph(str(val), styles_dict['ReportTableCell']) for val in mapper_func(item_obj)])
            
            data_table = Table(table_data, colWidths=col_widths)
            style = get_base_table_style() # Get base style
            apply_zebra_striping(style, len(table_data) -1) # Apply zebra striping
            data_table.setStyle(TableStyle(style)) # Apply the combined style
            story_list.append(data_table)
        else:
            story_list.append(Paragraph(f"No {title.split(' (')[0].lower().replace('handled','').replace('activity','').replace('registered','').strip()} for this period.", styles_dict['NoDataMessage']))
        story_list.append(Spacer(1, 0.15*inch))

    # Regular Events
    add_pdf_table_section(
        f"Events Handled ({regular_events.count()})",
        ["Ref #", "Client", "Event Name", "Location", "Start", "End", "Status"],
        regular_events,
        lambda e: [
            e.reference_number or '-', 
            f"{e.client.name if e.client else ''} {('('+e.client.company_name+')') if e.client and e.client.company_name else ''}".strip(),
            e.event_name or '-',
            e.event_location or '-', 
            timezone.localtime(e.start_date).strftime('%d/%m %H:%M') if e.start_date else '-',
            timezone.localtime(e.end_date).strftime('%d/%m %H:%M') if e.end_date else '-',
            e.get_status_display() or '-'
        ],
        [0.7*inch, 1.5*inch, 1.3*inch, 1.1*inch, 0.8*inch, 0.8*inch, 0.8*inch], # Adjusted widths
        story, styles
    )

    # Logistics Services Section
    add_pdf_table_section(
        f"Logistics Services Handled ({logistics_services.count()})",
        ["Ref #", "Client", "Service Title", "Primary Site", "Start", "End", "Status"],
        logistics_services,
        lambda s: [
            s.reference_number or '-', 
            f"{s.client.name if s.client else ''} {('('+s.client.company_name+')') if s.client and s.client.company_name else ''}".strip(),
            s.event_name or '-', 
            s.event_location or '-', 
            timezone.localtime(s.start_date).strftime('%d/%m %H:%M') if s.start_date else '-',
            timezone.localtime(s.end_date).strftime('%d/%m %H:%M') if s.end_date else '-',
            s.get_status_display() or '-'
        ],
        [0.7*inch, 1.5*inch, 1.3*inch, 1.1*inch, 0.8*inch, 0.8*inch, 0.8*inch],
        story, styles
    )
    
    # Rentals Section
    add_pdf_table_section(
        f"Rentals Handled ({rentals.count()})",
        ["Ref #", "Client", "Pickup", "Return", "Status"],
        rentals,
        lambda r: [
            r.reference_number or '-', 
            f"{r.client.name if r.client else ''} {('('+r.client.company_name+')') if r.client and r.client.company_name else ''}".strip(),
            timezone.localtime(r.start_date).strftime('%d/%m %H:%M') if r.start_date else '-',
            timezone.localtime(r.end_date).strftime('%d/%m %H:%M') if r.end_date else '-',
            r.get_status_display() or '-'
        ],
        [0.8*inch, 2.5*inch, 1.1*inch, 1.1*inch, 1.0*inch], # Adjusted widths
        story, styles
    )
    
    # Potential Page Break
    if len(story) > 30: # Heuristic for page break, adjust as necessary
        story.append(PageBreak())
        if logo_el: # Re-add header on new page if logo exists
            header_table_data_cont = [[header_left_elements, logo_el]]
            header_table_cont = Table(header_table_data_cont, colWidths=[5.5*inch, 1.5*inch])
            header_table_cont.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'), ('ALIGN', (0,0), (0,0), 'LEFT'), 
                ('ALIGN', (1,0), (1,0), 'RIGHT'), ('LEFTPADDING', (0,0), (0,0), 0), 
                ('RIGHTPADDING', (1,0), (1,0), 0),
            ]))
            story.append(header_table_cont)
        else:
            story.extend(header_left_elements) # Re-add text header
        story.append(Spacer(1, 0.1*inch))


    # Client Summary
    new_clients_qs = clients_summary.get('new_clients', Client.objects.none())
    add_pdf_table_section(
        f"New Clients Registered ({clients_summary.get('new_clients_count', 0)})",
        ["Client Name", "Company", "Email", "Registered On"],
        new_clients_qs,
        lambda c: [
            c.name or '-', 
            c.company_name or '-', 
            c.email or '-', 
            timezone.localtime(c.created_at).strftime('%Y-%m-%d') if c.created_at else '-'
        ],
        [1.75*inch, 1.75*inch, 2.0*inch, 1.5*inch],
        story, styles
    )

    # Item Usage Summary
    # For item usage, numbers should be right-aligned
    story.append(Paragraph(f"Item Usage Summary ({len(item_usage_summary)} item types used)", h2_style_pdf))
    if item_usage_summary:
        item_usage_headers = [Paragraph(h, table_header_cell_style_pdf) for h in ["Item Name", "SKU", "Times Used", "Total Qty Used"]]
        item_usage_data_pdf = [item_usage_headers]
        for item_data_dict in item_usage_summary:
            item_usage_data_pdf.append([
                Paragraph(str(item_data_dict.get('item__name', '-'))[:40], table_cell_style_pdf), # Truncate long names
                Paragraph(item_data_dict.get('item__sku', '-'), styles['ReportTableCellCenter']), # Center SKU
                Paragraph(str(item_data_dict.get('times_used', 0)), styles['ReportTableCellRight']), # Right align
                Paragraph(str(item_data_dict.get('total_quantity_used', 0)), styles['ReportTableCellRight']) # Right align
            ])
        
        item_usage_table = Table(item_usage_data_pdf, colWidths=[3.0*inch, 1.5*inch, 1.25*inch, 1.25*inch])
        style = get_base_table_style() # Get base style
        apply_zebra_striping(style, len(item_usage_data_pdf)-1) # Apply zebra striping
        # Additional alignment for numeric columns
        style.append(('ALIGN', (2,1), (3,-1), 'RIGHT')) 
        item_usage_table.setStyle(TableStyle(style))
        story.append(item_usage_table)
    else:
        story.append(Paragraph("No items were used in bookings this month.", no_data_style))

    doc.build(story)
    buffer.seek(0)
    # month_name_short_pdf = calendar.month_abbr[month_int]
    if str(month).isdigit():
        month_name_short_pdf = calendar.month_abbr[int(month)]
    else:
        month_name_short_pdf = "All"
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="clavis_monthly_summary_{month_name_short_pdf}_{year}.pdf"'
    return response


# --- Word Generation ---
# MODIFIED: Added logistics_services parameter
def generate_monthly_summary_docx(year, month, regular_events, logistics_services, rentals, clients_summary, item_usage_summary):
    document = Document()
    if str(month).isdigit():
        month_int = int(month)
        month_name_short_docx = calendar.month_abbr[month_int]
    else:
        month_int = None
        month_name_short_docx = "All"

    month_year_str = get_month_year_str(year, month)


    # Document Setup (Margins)
    section = document.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Title
    title_p = document.add_heading(f"Clavis Monthly Summary Report", level=0) # Level 0 for main title
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle_p = document.add_paragraph()
    subtitle_run = subtitle_p.add_run(f"Period: {month_year_str}")
    subtitle_run.italic = True
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph() # For spacing

    def add_table_to_docx_custom(doc, title, headers, data_queryset, data_mapper_func, numeric_cols_indices=None): # Added numeric_cols_indices
        doc.add_heading(title, level=2)
        
        has_data = False
        if hasattr(data_queryset, 'exists'): # For QuerySets
            has_data = data_queryset.exists()
        elif isinstance(data_queryset, list): # For lists
            has_data = bool(data_queryset)

        if has_data:
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            table.autofit = True # Let Word handle autofit for better spacing

            hdr_cells = table.rows[0].cells
            for i, header_text in enumerate(headers):
                cell_paragraph = hdr_cells[i].paragraphs[0]
                run = cell_paragraph.add_run(header_text)
                run.font.bold = True
                run.font.size = Pt(10)
                cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hdr_cells[i].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER


            for item_obj in data_queryset:
                row_data = data_mapper_func(item_obj)
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(row_data):
                    text_to_add = str(cell_text if cell_text is not None else '-')
                    cell_run = row_cells[i].paragraphs[0].add_run(text_to_add)
                    cell_run.font.size = Pt(9)
                    if numeric_cols_indices and i in numeric_cols_indices:
                        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph() 
        else:
            doc.add_paragraph(f"No {title.split(' (')[0].lower().replace('handled', '').replace('activity', '').replace('registered','').strip()} for this period.")
            doc.add_paragraph()

    # Regular Events
    add_table_to_docx_custom(document, f"Events Handled ({regular_events.count()})", 
                      ["Ref #", "Client", "Event Name", "Location", "Start", "End", "Status"], 
                      regular_events, 
                      lambda e: [e.reference_number or '-', 
                                 f"{e.client.name if e.client else ''} {('('+e.client.company_name+')') if e.client and e.client.company_name else ''}".strip(),
                                 e.event_name or '-',
                                 e.event_location or '-', 
                                 timezone.localtime(e.start_date).strftime('%Y-%m-%d %H:%M') if e.start_date else '-', 
                                 timezone.localtime(e.end_date).strftime('%Y-%m-%d %H:%M') if e.end_date else '-', 
                                 e.get_status_display() or '-'])

    # Logistics Services
    add_table_to_docx_custom(document, f"Logistics Services Handled ({logistics_services.count()})", 
                      ["Ref #", "Client", "Service Title", "Primary Site", "Start", "End", "Status"], 
                      logistics_services, 
                      lambda s: [s.reference_number or '-', 
                                 f"{s.client.name if s.client else ''} {('('+s.client.company_name+')') if s.client and s.client.company_name else ''}".strip(),
                                 s.event_name or '-', 
                                 s.event_location or '-', 
                                 timezone.localtime(s.start_date).strftime('%Y-%m-%d %H:%M') if s.start_date else '-', 
                                 timezone.localtime(s.end_date).strftime('%Y-%m-%d %H:%M') if s.end_date else '-', 
                                 s.get_status_display() or '-'])

    # Rentals
    add_table_to_docx_custom(document, f"Rentals Handled ({rentals.count()})", 
                      ["Ref #", "Client", "Pickup", "Return", "Status"], 
                      rentals, 
                      lambda r: [r.reference_number or '-', 
                                 f"{r.client.name if r.client else ''} {('('+r.client.company_name+')') if r.client and r.client.company_name else ''}".strip(),
                                 timezone.localtime(r.start_date).strftime('%Y-%m-%d %H:%M') if r.start_date else '-', 
                                 timezone.localtime(r.end_date).strftime('%Y-%m-%d %H:%M') if r.end_date else '-', 
                                 r.get_status_display() or '-'])
    
    # Client Summary
    new_clients_qs_docx = clients_summary.get('new_clients', Client.objects.none()) 
    add_table_to_docx_custom(document, f"New Clients Registered ({clients_summary.get('new_clients_count',0)})", 
                      ["Client Name", "Company", "Email", "Registered On"], 
                      new_clients_qs_docx, 
                      lambda c: [c.name or '-', 
                                 c.company_name or '-', 
                                 c.email or '-', 
                                 timezone.localtime(c.created_at).strftime('%Y-%m-%d') if c.created_at else '-'])
    
    # Item Usage Summary
    add_table_to_docx_custom(document, f"Item Usage Summary ({len(item_usage_summary)} item types used)",
                      ["Item Name", "SKU", "Times Used", "Total Qty Used"],
                      item_usage_summary,
                      lambda i: [
                          i.get('item__name', '-'), i.get('item__sku', '-'),
                          i.get('times_used', 0), i.get('total_quantity_used', 0)
                      ],
                      numeric_cols_indices=[2, 3]) # Times Used and Total Qty Used are numeric

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="clavis_monthly_summary_{month_name_short_docx}_{year}.docx"'
    return response